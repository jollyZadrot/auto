import os
import json
import re
import logging
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_LINE_SPACING, WD_COLOR_INDEX
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("docx_handler")


class DocxHandler:
    def __init__(self, client_data: dict, const_path: str = "const.json"):
        self.client_data = client_data
        self.const_path = const_path
        self.const_data = self._load_json(const_path) if os.path.exists(const_path) else {"web": "creditkasa.com.ua"}

    def _load_json(self, path: str) -> dict:
        try:
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception as e:
            logger.warning(f"Не вдалося завантажити {path}: {e}")
            return {}

    def _clean_and_paint_black(self, doc):
        header_pattern = re.compile(r"^[IІVХX]+\.\s+[А-ЯЩЬЮЯЄІЇҐA-Z]", re.IGNORECASE)
        for paragraph in doc.paragraphs:
            p_text = paragraph.text.strip()
            is_header = bool(header_pattern.match(p_text))
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.highlight_color = WD_COLOR_INDEX.GRAY_25 if is_header else None

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        p_text = paragraph.text.strip()
                        is_header = bool(header_pattern.match(p_text))
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            run.font.highlight_color = WD_COLOR_INDEX.GRAY_25 if is_header else None

    def _format_rate_comma(self, rate_str: str) -> str:
        try:
            val = float(rate_str.replace(',', '.'))
            return f"{val:.2f}".replace('.', ',') + " %"
        except ValueError:
            return f"{rate_str} %"

    def _convert_rate_to_words(self, rate_str: str) -> tuple:
        try:
            val = float(rate_str.replace(",", "."))
            if val == 1.0 or val == 1:
                return "1.00%", "одна ціла, нуль сотих процентів"
            elif val == 0.74:
                return "0.74%", "нуль цілих, сімдесят чотири сотих процентів"
            elif val == 1.5:
                return "1.50%", "одна ціла, п'ятдесят сотих процентів"

            parts = f"{val:.2f}".split('.')
            whole, frac = int(parts[0]), int(parts[1])
            num_words = {0: "нуль", 1: "одна", 2: "дві", 3: "три", 4: "чотири", 5: "п'ять", 6: "шість", 7: "сім",
                         8: "вісім", 9: "дев'ять"}
            frac_words = {0: "нуль", 50: "п'ятдесят", 74: "сімдесят чотири"}

            w_txt = num_words.get(whole, str(whole))
            f_txt = frac_words.get(frac, str(frac))
            w_suf = "ціла" if whole == 1 else "цілих"
            return f"{val:.2f}%", f"{w_txt} {w_suf}, {f_txt} сотих процентів"
        except Exception:
            return f"{rate_str}%", "вказана кількість процентів"

    def _replace_in_p(self, paragraph, pattern_str, replacement):
        pattern = re.compile(pattern_str, re.IGNORECASE)
        replaced = False
        for run in paragraph.runs:
            if pattern.search(run.text):
                run.text = pattern.sub(replacement, run.text)
                replaced = True
        if not replaced:
            full_txt = "".join(r.text for r in paragraph.runs)
            if pattern.search(full_txt):
                fmts = [(r.font.name, r.font.size, r.bold, r.italic) for r in paragraph.runs]
                paragraph.text = pattern.sub(replacement, full_txt)
                if paragraph.runs and fmts:
                    paragraph.runs[0].font.name, paragraph.runs[0].font.size = fmts[0][0], fmts[0][1]
                    paragraph.runs[0].bold, paragraph.runs[0].italic = fmts[0][2], fmts[0][3]

    def _remove_p(self, paragraph):
        element = paragraph._element
        if element.getparent() is not None:
            element.getparent().remove(element)

    def _insert_cloned_p_before(self, target_p, text, clone_from=None):
        ref_p = clone_from if clone_from else target_p
        new_p = target_p.insert_paragraph_before()
        if ref_p.paragraph_format.left_indent is not None:
            new_p.paragraph_format.left_indent = ref_p.paragraph_format.left_indent
        if ref_p.paragraph_format.first_line_indent is not None:
            new_p.paragraph_format.first_line_indent = ref_p.paragraph_format.first_line_indent
        match = re.match(r"^(\s*[-*•]\s*)", ref_p.text)
        prefix = match.group(1) if match else ""
        final_text = f"{prefix}{text}"
        r = new_p.add_run(final_text)
        r.bold, r.italic = False, False
        return new_p

    def _insert_p_aligned(self, p_target, p_ref, text, insert_before=False):
        """Ідеальне математичне вирівнювання по лінії маркерів (червоній лінії)."""
        p_ref = p_ref if p_ref else p_target
        new_p_xml = OxmlElement('w:p')

        if insert_before:
            p_target._p.addprevious(new_p_xml)
        else:
            p_target._p.addnext(new_p_xml)

        new_p = Paragraph(new_p_xml, p_target._parent)

        # Вираховуємо точну позицію дефіса (буліта)
        left_ind = p_ref.paragraph_format.left_indent
        first_line_ind = p_ref.paragraph_format.first_line_indent
        val_left = left_ind.twips if left_ind is not None else 0
        val_first = first_line_ind.twips if first_line_ind is not None else 0

        bullet_pos = val_left + val_first

        # Ставимо текст рівно по цій лінії
        new_p.paragraph_format.left_indent = Pt(bullet_pos / 20.0) if bullet_pos else Pt(0)
        new_p.paragraph_format.first_line_indent = Pt(0)

        new_p.paragraph_format.space_before = Pt(0)
        new_p.paragraph_format.space_after = Pt(0)
        new_p.paragraph_format.line_spacing = 1.0
        new_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        r = new_p.add_run(text)
        r.bold, r.italic = False, False
        return new_p

    def process_file(self, input_path: str, output_path: str) -> str:
        logger.info(f"Обробка DOCX: {input_path}")
        doc = Document(input_path)

        web_domain = self.const_data.get("web", "creditkasa.com.ua")
        is_target_domain = False
        for p in doc.paragraphs:
            if web_domain.lower() in p.text.lower():
                is_target_domain = True
                break

        for paragraph in doc.paragraphs:
            self._replace_in_p(paragraph, r"creditkasa\.(com\.)?ua", web_domain)

        otp = self.client_data.get("otp", "")
        if otp:
            for paragraph in doc.paragraphs:
                if "на виконання зазначених вимог" in paragraph.text.lower() or "ідентифікатор" in paragraph.text.lower():
                    self._replace_in_p(paragraph, r"(ідентифікатор\s+)[A-Za-zА-Яа-я0-9іІїЇєЄґҐ]+", rf"\g<1>{otp}")

        fee_val = str(self.client_data.get("fee", ""))
        std_val = str(self.client_data.get("standard_rate", ""))
        rate_181_val = str(self.client_data.get("rate_181", ""))
        rate_181_words = str(self.client_data.get("rate_181_words", ""))
        red_val = str(self.client_data.get("reduced_rate", ""))
        promo_val = str(self.client_data.get("promo_rate", ""))
        pref_val = str(self.client_data.get("preferential_rate", ""))

        start_idx, end_idx = -1, -1
        for i, p in enumerate(doc.paragraphs):
            txt = p.text.lower()
            if "на наступних умовах:" in txt:
                start_idx = i + 1
            elif start_idx != -1 and "* базовий період" in txt and "проміжки часу" in txt:
                end_idx = i
                break

        if start_idx != -1 and end_idx != -1:
            cond_paragraphs = doc.paragraphs[start_idx:end_idx]

            p_fee = p_red = p_promo = p_pref = p_std_title = None
            p_std_subs = []

            for p in cond_paragraphs:
                txt = p.text.lower()
                if p_std_title is not None:
                    p_std_subs.append(p)
                elif "комісія за видачу" in txt:
                    p_fee = p
                elif "промо-ставка" in txt or "(промо-ставка)" in txt:
                    p_promo = p
                elif "пільгова" in txt:
                    p_pref = p
                elif "знижена" in txt:
                    p_red = p
                elif "стандартна % ставка" in txt:
                    p_std_title = p

            p_bullet_ref = p_fee if p_fee else (p_red if p_red else None)

            if fee_val and fee_val not in ["0", "0.00", ""]:
                if p_fee: self._replace_in_p(p_fee, r"\d+(?:[.,]\d+)?\s*%", self._format_rate_comma(fee_val))
            elif p_fee:
                self._remove_p(p_fee)

            if red_val and red_val not in ["0", "0.00", ""]:
                if p_red: self._replace_in_p(p_red, r"\d+(?:[.,]\d+)?\s*%", self._format_rate_comma(red_val))
            elif p_red:
                self._remove_p(p_red)

            if promo_val and promo_val not in ["0", "0.00", ""]:
                if p_promo:
                    self._replace_in_p(p_promo, r"\d+(?:[.,]\d+)?\s*%", self._format_rate_comma(promo_val))
                elif p_std_title and p_bullet_ref:
                    self._insert_cloned_p_before(p_std_title,
                                                 f"знижена % ставка (Промо-ставка) – {self._format_rate_comma(promo_val)} в день;",
                                                 clone_from=p_bullet_ref)
            elif p_promo:
                self._remove_p(p_promo)

            if pref_val and pref_val not in ["0", "0.00", ""]:
                if p_pref:
                    self._replace_in_p(p_pref, r"\d+(?:[.,]\d+)?\s*%", self._format_rate_comma(pref_val))
                elif p_std_title and p_bullet_ref:
                    self._insert_cloned_p_before(p_std_title,
                                                 f"пільгова % ставка – {self._format_rate_comma(pref_val)} в день;",
                                                 clone_from=p_bullet_ref)
            elif p_pref:
                self._remove_p(p_pref)

            if p_std_title and p_bullet_ref:
                if is_target_domain:
                    # Вставляємо "стандартна % ставка:" ідеально рівно
                    new_std_title = self._insert_p_aligned(p_std_title, p_bullet_ref, "стандартна % ставка:",
                                                           insert_before=True)
                    self._remove_p(p_std_title)
                    for p in p_std_subs: self._remove_p(p)

                    std_pct, std_words = self._convert_rate_to_words(std_val)

                    txt2 = f"*-{rate_181_val}% ({rate_181_words}) за кожен день користування Кредитом, яка застосовується у період починаючи з 181 (ста вісімдесят першого) календарного дня дії Договору і до закінчення строку дії цього Договору або до дати фактичного повернення всієї суми Кредиту (до тієї із зазначених дат, яка настане раніше)."
                    txt1 = f"*- {std_pct} ({std_words}) за кожен день користування Кредитом, яка застосовується протягом перших 180 (ста вісімдесяти) календарних днів з дати укладення цього Договору. В цей період можливе використання Позичальником права користування Кредитом за Промо-ставкою та/або Зниженою та/або Пільговою процентною ставкою."

                    # Вставляємо текст про 181 і 180 днів
                    self._insert_p_aligned(new_std_title, p_bullet_ref, txt2, insert_before=False)
                    self._insert_p_aligned(new_std_title, p_bullet_ref, txt1, insert_before=False)
                else:
                    self._insert_cloned_p_before(p_std_title,
                                                 f"стандартна % ставка – {self._format_rate_comma(std_val)} в день;",
                                                 clone_from=p_bullet_ref)
                    self._remove_p(p_std_title)
                    for p in p_std_subs: self._remove_p(p)

        self._clean_and_paint_black(doc)
        doc.save(output_path)
        logger.info(f"Оброблений файл збережено: {output_path}")
        return output_path
